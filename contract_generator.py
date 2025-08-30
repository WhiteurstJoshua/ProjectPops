
"""
Powell CM Solutions - Contract Generator (Enhanced Rebuild v2)
--------------------------------------------------------------
Generates a customized Master Services Agreement + Work Order in Word (.docx).

Requirements:
- Python 3.x
- python-docx  (pip install python-docx)

Usage:
- Run:  python contract_generator.py
- Answer prompts (press Enter to accept defaults).
- Output: Generated_Agreement_YYYYMMDD_HHMM.docx
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# ---------------------------- Helpers ----------------------------

def add_paragraph(doc, text, size=11):
    p = doc.add_paragraph(text)
    if p.runs:
        p.runs[0].font.name = 'Calibri'
        p.runs[0].font.size = Pt(size)
    return p

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_table_rate_card(doc, rates):
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Role'
    hdr_cells[1].text = 'Hourly Rate (USD)'
    for role, rate in rates.items():
        row_cells = table.add_row().cells
        row_cells[0].text = role
        row_cells[1].text = f"${float(rate):,.2f}"
    return table

def menu_choice(prompt, options, default_idx=0):
    print(prompt)
    for i, opt in enumerate(options, 1):
        print(f"  {i}) {opt}")
    raw = input(f"Select [default {default_idx+1}]: ").strip()
    if not raw:
        return options[default_idx]
    try:
        idx = int(raw) - 1
        if 0 <= idx < len(options):
            return options[idx]
    except:
        pass
    return options[default_idx]

def prompt_text(prompt, default=""):
    raw = input(f"{prompt} [{default}]: ").strip()
    return raw if raw else default

def prompt_money(prompt, default="0"):
    raw = input(f"{prompt} [{default}]: ").replace(",", "").strip()
    try:
        return float(raw) if raw else float(default or 0)
    except:
        return float(default or 0)

def timestamp_suffix():
    return datetime.now().strftime("%Y%m%d_%H%M")

# ---------------------------- Builder ----------------------------

def build_agreement(data):
    doc = Document()

    # Title
    h = doc.add_heading('Master Professional Services Agreement (Generated)', 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Parties / Header
    add_paragraph(doc, f"This Agreement is entered into as of {data['effective_date']} by and between "
                       f"{data['client_name']} (\"Client\") and {data['firm_name']} (\"Consultant\").")
    add_paragraph(doc, "Client and Consultant are together the “Parties.”")

    # Work Order Summary
    add_heading(doc, 'Work Order Summary', level=1)
    add_paragraph(doc, f"Project: {data['project_name']}")
    add_paragraph(doc, f"Role: {data['role']}")
    add_paragraph(doc, f"Relationship: {data['relationship']}")
    add_paragraph(doc, f"Prime Contract Reference (if Sub): {data['prime_reference']}")
    add_paragraph(doc, f"Term: {data['term']}")

    # Recitals
    add_heading(doc, 'Recitals', level=1)
    add_paragraph(doc, "A. Client desires to engage Consultant to provide program management, construction management "
                       "(as advisor), owner’s representation, developer advisory, and/or related professional consulting services.")
    add_paragraph(doc, "B. Consultant is duly qualified and willing to perform such services under the terms of this Agreement.")

    # Articles
    add_heading(doc, 'Article 1 – Master Engagement & Work Orders', level=1)
    add_paragraph(doc, "1.1 Master Agreement. This Agreement sets general terms for all services. Specific scope, fees, "
                       "schedules, and special terms will be set forth in written Work Orders executed by the Parties and incorporated herein.")
    add_paragraph(doc, "1.2 Prime/Sub Flexibility. Consultant may perform services as prime directly for Client or as a subconsultant, "
                       "as indicated in the Work Order. Flow‑down obligations from any prime contract apply to Consultant only to the extent "
                       "expressly identified in the Work Order.")
    add_paragraph(doc, "1.3 Independent Contractor; No Authority to Bind. Consultant is an independent contractor and shall not bind Client without written authority.")

    add_heading(doc, 'Article 2 – Standard of Care; Personnel', level=1)
    add_paragraph(doc, "2.1 Standard of Care. Consultant shall perform services with the care and skill ordinarily used by similar professionals "
                       "practicing under similar conditions at the same time and locality.")
    add_paragraph(doc, "2.2 Key Personnel. If key personnel are identified, Consultant shall not reassign them without reasonable notice and suitable replacement.")

    add_heading(doc, 'Article 3 – Compensation & Payment', level=1)
    if data['compensation'] == "Hourly":
        add_paragraph(doc, f"3.1 Fees. Hourly per Rate Exhibit E with a Not‑to‑Exceed amount of ${data['nte']:,.2f} without prior written approval.")
    elif data['compensation'] == "Lump Sum":
        add_paragraph(doc, f"3.1 Fees. Lump Sum fee of ${data['lump_sum']:,.2f}, payable per milestones set forth in Exhibit B.")
    else:
        add_paragraph(doc, f"3.1 Fees. Hybrid: Hourly per Exhibit E with a monthly cap of ${data['monthly_cap']:,.2f}.")
    add_paragraph(doc, "3.2 Reimbursable Expenses. Billed at actual cost per Exhibit D unless otherwise stated.")
    add_paragraph(doc, f"3.3 Invoices & Payment. Invoices monthly; payment due net {data['net_days']}. "
                       "Overdue balances accrue interest at 1% per month or the maximum allowed by law.")

    add_heading(doc, 'Article 4 – Insurance', level=1)
    if data['insurance'] == "Standard":
        add_paragraph(doc, "GL $1M each / $2M aggregate; Auto $1M CSL; WC Statutory; Employers $500k; Professional Liability $2M aggregate.")
    elif data['insurance'] == "Expanded":
        add_paragraph(doc, "GL $2M each / $4M aggregate; Auto $1M CSL; WC Statutory; Employers $1M; Professional Liability $5M aggregate.")
    else:
        add_paragraph(doc, "GL $1M each; Auto N/A if no driving; WC Statutory; Employers $500k; Professional Liability $1M aggregate.")
    if data['ai_required']:
        add_paragraph(doc, "Additional Insured status will be provided where required by the Work Order or prime contract, to the extent commercially available.")

    add_heading(doc, 'Article 5 – Ownership; License; Confidentiality', level=1)
    if data['ip_assignment'] == "License":
        add_paragraph(doc, "5.1 Instruments of Service. Upon full payment, Client receives a non‑exclusive license to use deliverables for the Project identified in the Work Order. Consultant retains IP rights.")
    else:
        add_paragraph(doc, "5.1 Instruments of Service. Upon full payment, Consultant assigns to Client the ownership of deliverables for the Project identified in the Work Order (excluding Consultant’s pre‑existing tools).")
    add_paragraph(doc, "5.2 Confidentiality. Each Party shall keep in confidence non‑public information received from the other and use it solely for the Project.")
    if data['include_nda']:
        add_paragraph(doc, "5.3 Mutual NDA. The Parties agree not to disclose Confidential Information except to those with a need to know who are bound by confidentiality obligations; "
                           "to protect such information with at least the same degree of care as used to protect their own; and to return or destroy such information upon written request, "
                           "subject to legal and record‑keeping requirements.")

    add_heading(doc, 'Article 6 – Indemnification; Limitation of Liability', level=1)
    add_paragraph(doc, "6.1 Consultant Indemnity. To the extent caused by Consultant’s negligence, gross negligence, or willful misconduct, Consultant shall indemnify and hold harmless Client from third‑party claims for bodily injury, death, or tangible property damage. This indemnity excludes Client’s negligence.")
    add_paragraph(doc, "6.2 Client Indemnity. Client shall indemnify and hold harmless Consultant from third‑party claims to the extent caused by Client’s negligence or willful misconduct.")
    lol_text = f"{data['lol_multiplier']}x the fees paid for the applicable Work Order" if data['lol_multiplier'] else "two (2) times the fees paid"
    add_paragraph(doc, f"6.3 Limitation of Liability. Consultant’s aggregate liability under this Agreement and any Work Order shall not exceed {lol_text}. "
                       "Neither Party shall be liable for consequential, incidental, or special damages.")
    if data['include_dei']:
        add_paragraph(doc, "6.4 Inclusion & Non‑Discrimination. Consultant shall endeavor to utilize a diverse workforce and comply with applicable non‑discrimination laws and Client’s reasonable inclusion objectives.")

    add_heading(doc, 'Article 7 – Changes; Suspension; Termination', level=1)
    add_paragraph(doc, "7.1 Changes require written authorization via amendment to the Work Order.")
    add_paragraph(doc, "7.2 Suspension. Client may suspend upon written notice; Consultant shall be paid for work performed and reasonable demobilization/remobilization costs.")
    add_paragraph(doc, f"7.3 Termination for Convenience. Either Party may terminate a Work Order on {data['termination_notice_days']} days’ written notice. Consultant shall be paid for services performed and costs incurred through termination.")

    add_heading(doc, 'Article 8 – Dispute Resolution', level=1)
    if data['dispute'] == "Litigation":
        add_paragraph(doc, f"Disputes shall be resolved in the state courts of {data['venue_county']}, {data['venue_state']}. Jury trial waived to the extent permitted by law.")
    elif data['dispute'] == "Arbitration":
        add_paragraph(doc, f"Disputes shall be mediated first; if unresolved, finally resolved by binding arbitration under the AAA Construction Industry Rules. Seat: {data['venue_city']}, {data['venue_state']}.")
    else:
        add_paragraph(doc, f"Disputes shall be mediated first; if unresolved, litigated in the state courts of {data['venue_county']}, {data['venue_state']}.")

    add_heading(doc, 'Article 9 – Miscellaneous', level=1)
    add_paragraph(doc, f"9.1 Governing Law. The laws of {data['venue_state']} apply.")
    add_paragraph(doc, "9.2 Assignment. Neither Party may assign without written consent, except to affiliates in connection with a merger, acquisition, or reorganization.")
    add_paragraph(doc, "9.3 Entire Agreement. This Agreement, together with applicable Work Orders and Exhibits, constitutes the entire agreement between the Parties.")

    # Scope Library
    add_heading(doc, 'Exhibit A – Scope of Services (Role-Based Library)', level=1)
    scope_map = {
        "Owner’s Representative": [
            "Design phase coordination; value analysis; constructability; permitting roadmap.",
            "Procurement support (RFPs, bid leveling), recommendations; contract administration support.",
            "Construction monitoring; pay app/change order review; schedule analysis; punch/turnover oversight."
        ],
        "Program Manager": [
            "PMO governance; executive dashboards; stage‑gate reviews; RAID/risk tracking.",
            "Master schedule (L1–L3); document control; cost/schedule reporting; baseline & forecasts."
        ],
        "Construction Manager (Advisor)": [
            "Preconstruction estimating; budget/schedule alignment; logistics planning.",
            "Submittal/RFI workflow; reporting cadence; change management support; claims avoidance."
        ],
        "Developer Advisory": [
            "Feasibility and entitlement support; utilities coordination; community engagement planning.",
            "Pro forma inputs; delivery strategy; risk register and mitigation planning; lender/partner reporting."
        ],
        "Subconsultant": [
            "Discipline‑specific tasks aligned with prime contract flow‑downs.",
            "Coordinate deliverables and schedule under prime consultant’s direction."
        ]
    }
    chosen = scope_map.get(data['role'], ["Custom scope to be attached."])
    for item in chosen:
        add_paragraph(doc, "• " + item)

    # Compensation
    add_heading(doc, 'Exhibit B – Compensation', level=1)
    if data['compensation'] == "Hourly":
        add_paragraph(doc, f"Hourly per Exhibit E; Not‑to‑Exceed ${data['nte']:,.2f} without prior written approval.")
    elif data['compensation'] == "Lump Sum":
        add_paragraph(doc, f"Lump Sum Fee: ${data['lump_sum']:,.2f}, payable per agreed milestones.")
    else:
        add_paragraph(doc, f"Hybrid: Hourly per Exhibit E with monthly cap ${data['monthly_cap']:,.2f}.")

    # Insurance
    add_heading(doc, 'Exhibit C – Insurance', level=1)
    if data['insurance'] == "Standard":
        add_paragraph(doc, "GL $1M each / $2M agg; Auto $1M CSL; WC Statutory; Employers $500k; Professional $2M agg.")
    elif data['insurance'] == "Expanded":
        add_paragraph(doc, "GL $2M each / $4M agg; Auto $1M CSL; WC Statutory; Employers $1M; Professional $5M agg.")
    else:
        add_paragraph(doc, "GL $1M each; Auto N/A if no driving; WC Statutory; Employers $500k; Professional $1M agg.")

    # Reimbursables
    add_heading(doc, 'Exhibit D – Reimbursable Expenses', level=1)
    add_paragraph(doc, "Travel (coach airfare), lodging at GSA per diem, mileage at IRS rate, meals per diem, printing/ repro, permits/fees, courier/delivery, "
                       "pre‑approved software/hosting, and meeting/event costs. Billed at actual cost, no markup.")

    # Rate card
    add_heading(doc, 'Exhibit E – Rate Schedule', level=1)
    add_paragraph(doc, "Standard rates (edit in prompts or here):")
    add_table_rate_card(doc, data['rates'])
    add_paragraph(doc, f"Annual adjustment: up to {data['annual_increase_cap']}% with thirty (30) days’ notice, unless otherwise agreed.")

    # Disputes
    add_heading(doc, 'Exhibit F – Dispute Resolution', level=1)
    if data['dispute'] == "Litigation":
        add_paragraph(doc, f"Exclusive venue and jurisdiction: state courts of {data['venue_county']}, {data['venue_state']}.")
    elif data['dispute'] == "Arbitration":
        add_paragraph(doc, f"Mediation first; if unresolved, binding arbitration (AAA Construction Industry Rules). Seat: {data['venue_city']}, {data['venue_state']}.")
    else:
        add_paragraph(doc, f"Mediation first; if unresolved, litigation in state courts of {data['venue_county']}, {data['venue_state']}.")

    # Signature Blocks
    doc.add_page_break()
    add_heading(doc, 'Signatures', level=1)
    add_paragraph(doc, f"{data['client_name']}", size=12)
    add_paragraph(doc, "By: _______________________________")
    add_paragraph(doc, "Name: _____________________________")
    add_paragraph(doc, "Title: ______________________________")
    add_paragraph(doc, "Date: ______________________________")
    add_paragraph(doc, "")
    add_paragraph(doc, f"{data['firm_name']}", size=12)
    add_paragraph(doc, "By: _______________________________")
    add_paragraph(doc, "Name: _____________________________")
    add_paragraph(doc, "Title: ______________________________")
    add_paragraph(doc, "Date: ______________________________")

    out_path = f"Generated_Agreement_{timestamp_suffix()}.docx"
    doc.save(out_path)
    return out_path

# ---------------------------- Main ----------------------------

def main():
    print("=== Powell CM Solutions - Contract Generator (Enhanced v2) ===")
    # Defaults
    defaults = {
        "firm_name": "Powell CM Solutions, LLC",
        "venue_county": "Cook County",
        "venue_state": "Illinois",
        "venue_city": "Chicago",
        "net_days": 30,
        "termination_notice_days": 15,
        "lol_multiplier": 2,
        "annual_increase_cap": 4,
        "ip_assignment": "License",  # or "Assignment"
        "include_nda": True,
        "include_dei": True,
        "ai_required": True,  # Additional Insured language
        "rates": {
            "Principal / Executive": 250.00,
            "Senior Project Manager": 200.00,
            "Project Manager": 175.00,
            "Project Engineer": 150.00,
            "Coordinator / Admin": 100.00
        }
    }

    # Parties & basics
    effective_date = prompt_text("Effective Date (e.g., August 29, 2025)", "")
    client_name = prompt_text("Client Legal Name", "")
    firm_name = prompt_text("Your Firm Legal Name", defaults["firm_name"])
    project_name = prompt_text("Project Name", "")
    term = prompt_text("Term (e.g., Sept 1, 2025 – June 30, 2026)", "")

    role = menu_choice("Select Role:", [
        "Owner’s Representative",
        "Program Manager",
        "Construction Manager (Advisor)",
        "Developer Advisory",
        "Subconsultant"
    ], default_idx=1)

    relationship = menu_choice("Select Relationship:", [
        "Prime",
        "Subconsultant"
    ], default_idx=0)

    prime_reference = ""
    if relationship == "Subconsultant":
        prime_reference = prompt_text("Prime Contract Reference / Project Flow-down Identifier", "[Attach Reference]")

    compensation = menu_choice("Compensation Method:", [
        "Hourly",
        "Lump Sum",
        "Hybrid"
    ], default_idx=0)

    nte = lump_sum = monthly_cap = 0.0
    if compensation == "Hourly":
        nte = prompt_money("If Hourly, Not‑to‑Exceed amount (number only)", "0")
    elif compensation == "Lump Sum":
        lump_sum = prompt_money("If Lump Sum, amount (number only)", "0")
    else:
        monthly_cap = prompt_money("If Hybrid, monthly cap (number only)", "0")

    insurance = menu_choice("Insurance Tier:", [
        "Standard",
        "Expanded",
        "Reduced"
    ], default_idx=0)

    dispute = menu_choice("Dispute Resolution:", [
        "Litigation",
        "Arbitration",
        "Mediation-then-Court"
    ], default_idx=0)

    venue_county = prompt_text("Venue County", defaults["venue_county"])
    venue_state = prompt_text("Venue State", defaults["venue_state"])
    venue_city = prompt_text("Venue City (for arbitration seat)", defaults["venue_city"])

    # Optional toggles
    print("\n--- Optional Clauses (Enter to keep default) ---")
    net_days = int(prompt_text("Payment Terms - Net Days", str(defaults["net_days"])) or defaults["net_days"])
    termination_notice_days = int(prompt_text("Termination Notice Days", str(defaults["termination_notice_days"])) or defaults["termination_notice_days"])
    lol_multiplier = int(prompt_text("Limitation of Liability multiplier (x fees)", str(defaults["lol_multiplier"])) or defaults["lol_multiplier"])
    annual_increase_cap = int(prompt_text("Annual Rate Increase Cap (%)", str(defaults["annual_increase_cap"])) or defaults["annual_increase_cap"])

    ip_assignment = menu_choice("Deliverables ownership:", ["License", "Assignment"], default_idx=0)
    include_nda = menu_choice("Include short mutual NDA?", ["Yes", "No"], default_idx=0) == "Yes"
    include_dei = menu_choice("Include Inclusion/Non-Discrimination statement?", ["Yes", "No"], default_idx=0) == "Yes"
    ai_required = menu_choice("Include Additional Insured language?", ["Yes", "No"], default_idx=0) == "Yes"

    # Rate Card (edit inline)
    print("\nCurrent Rate Card (press Enter to keep):")
    rates = {}
    for role_name, default_rate in defaults["rates"].items():
        raw = prompt_text(f"Rate for {role_name}", f"{default_rate}")
        try:
            rates[role_name] = float(raw)
        except:
            rates[role_name] = default_rate

    data = {
        "effective_date": effective_date or "____ ______, 20__",
        "client_name": client_name or "[Client Legal Name]",
        "firm_name": firm_name or defaults["firm_name"],
        "project_name": project_name or "[Project Name]",
        "term": term or "[Start – End]",
        "role": role,
        "relationship": relationship,
        "prime_reference": prime_reference,
        "compensation": compensation,
        "nte": nte,
        "lump_sum": lump_sum,
        "monthly_cap": monthly_cap,
        "insurance": insurance,
        "dispute": dispute,
        "venue_county": venue_county,
        "venue_state": venue_state,
        "venue_city": venue_city,
        "net_days": net_days,
        "termination_notice_days": termination_notice_days,
        "lol_multiplier": lol_multiplier,
        "annual_increase_cap": annual_increase_cap,
        "ip_assignment": ip_assignment,
        "include_nda": include_nda,
        "include_dei": include_dei,
        "ai_required": ai_required,
        "rates": rates
    }

    out = build_agreement(data)
    print(f"\nDone! Created: {out}")

if __name__ == "__main__":
    main()
